# Importing all the required Libraries.
from bs4 import BeautifulSoup
import requests
from selenium import webdriver
from docx import Document

#Creating an Document Object.
mydoc = Document()
chapters = []
chap_list = []

# For running the chrome driver.
chrome_driver_path = 'F:\Softwares\chromedriver.exe'
driver = webdriver.Chrome(executable_path=chrome_driver_path)

# URL of the Website you want to start scrapping on.
page_path = "https://sololeveling.com/manga/solo-leveling/solo-leveling-chapter-1/"
driver.get(page_path)

# For testing I used for loop.
# To get all the chapters use while loop with true as condition
# and when there are no more chapters return False so that it stops the program.

for i in range(10):
    
    # Using different methods of selenium to extract the Title of the Chapter.
    chap_title= driver.find_element_by_css_selector('.breadcrumb .active').text
    chapters.append(chap_title)
    chap_no = int(chap_title[22:])
    chap_list.append(chap_no)
    
    # Extracting the Summary of Chapter
    chap_summary = driver.find_element_by_css_selector('.reading-content .text-left').text
    
    # Storing all the data into a word document
    mydoc.add_heading(chap_title,0)
    mydoc.add_paragraph(chap_summary)
    mydoc.add_section()
    mydoc.save('SoloLeveling.docx')
    
    # Navigating to Next Chapter.
    # nav_next = driver.find_element_by_xpath('//*[@id="manga-reading-nav-head"]/div/div[3]/div/div/a')
    nav_next = driver.find_element_by_css_selector('.nav-next .next_page').get_attribute('href')
    print(nav_next)
    driver.get(nav_next)

# Finally saving the document. 
mydoc.save('SoloLeveling.docx')
print(chapters)
print(chap_list)

driver.close()  # Closes Tab
# # driver.quit()  # Closes Browser
